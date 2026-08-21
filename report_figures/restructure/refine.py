#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Tinh chỉnh cấu trúc Chương 2 sau khi gộp.

1. Chuyển 14 ảnh chụp màn hình đang nằm mồ côi trong mục thiết kế về đúng các
   chú thích hình đang thiếu ảnh ở phần hiện thực hóa.
2. Tách mục "Thiết kế giao diện và bảo đảm an toàn thông tin" thành hai mục.
3. Đánh số lại các mục phía sau và sửa tiêu đề thiếu số ở Chương 1.
"""
import sys, docx
from copy import deepcopy
from docx.shared import Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, DST = sys.argv[1], sys.argv[2]
d = docx.Document(SRC)
ps = d.paragraphs
A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
WP = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'


def set_par_text(p, text):
    runs = p.runs
    if not runs:
        p.add_run(text); return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def resize(par, width_in):
    """Đổi bề rộng ảnh trong đoạn, giữ nguyên tỉ lệ."""
    for inline in par._p.findall('.//' + WP + 'inline'):
        ext = inline.find(WP + 'extent')
        if ext is None:
            continue
        cx, cy = int(ext.get('cx')), int(ext.get('cy'))
        ncx = int(Inches(width_in))
        ncy = int(cy * ncx / cx)
        ext.set('cx', str(ncx)); ext.set('cy', str(ncy))
        for e in inline.findall('.//' + A + 'ext'):
            if e.get('cx') is not None:
                e.set('cx', str(ncx)); e.set('cy', str(ncy))


# ── 1. Xác định khối ảnh mồ côi (giữa tiêu đề mục 5 và tiểu mục 5.1) ────────
H2_UI = 288
H3_51 = 309
imgs, blanks = [], []
for i in range(H2_UI + 1, H3_51):
    (imgs if ps[i]._p.findall('.//' + A + 'blip') else blanks).append(i)
assert len(imgs) == 14, len(imgs)
print('ảnh mồ côi:', imgs)
print('đoạn trống kèm theo:', blanks)

# chú thích đích  ->  (danh sách ảnh theo thứ tự, bề rộng inch)
def cap(num):
    return next(i for i, p in enumerate(ps)
                if p.text.strip().startswith('Hình %s.' % num))

PLAN = [
    (cap('2.13'), [imgs[0], imgs[1], imgs[13]], 5.4),   # đăng nhập · trang chính (admin) · trang chính (người dùng)
    (cap('2.14'), [imgs[2], imgs[3]], 6.3),             # tải tài liệu · lựa chọn công cụ và kết quả
    (cap('2.15'), [imgs[9], imgs[10], imgs[11], imgs[12]], 4.6),  # tổng quan · người dùng · nhật ký · tệp tin
    (cap('2.18'), [imgs[4], imgs[5], imgs[6]], 5.4),    # thẻ Nguồn · Hình ảnh · JSON
    (cap('2.21'), [imgs[7], imgs[8]], 6.3),             # thư viện tài liệu · kết quả OCR đầy đủ
]
moved = 0
for cap_i, srcs, w in PLAN:
    anchor = ps[cap_i]._p
    for s in srcs:
        par = ps[s]
        resize(par, w)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.keep_with_next = True
        anchor.addprevious(par._p)      # lxml: chuyển phần tử
        moved += 1
print('đã chuyển %d ảnh về %d chú thích' % (moved, len(PLAN)))

# xóa các đoạn trống còn lại trong khối cũ
for i in blanks:
    ps[i]._p.getparent().remove(ps[i]._p)

# ── 2. Tách mục 5 thành hai mục ────────────────────────────────────────────
set_par_text(ps[H2_UI], '5. Thiết kế giao diện hệ thống')
new_h2 = deepcopy(ps[H2_UI]._p)          # sao chép để giữ nguyên định dạng Heading 2
ps[325]._p.addprevious(new_h2)
from docx.text.paragraph import Paragraph
set_par_text(Paragraph(new_h2, ps[H2_UI]._parent), '6. Thiết kế bảo mật và an toàn thông tin')

# ── 3. Đánh số lại các mục phía sau ────────────────────────────────────────
RENUM = {
    325: '6.1. Xác thực, phân quyền và kiểm soát truy cập',
    330: '6.2. Bảo vệ tài liệu và dữ liệu lưu trữ',
    336: '6.3. Lưu vết hoạt động, sao lưu và xử lý sự cố',
    346: '7. Môi trường, công nghệ và công cụ xây dựng hệ thống',
    347: '7.1. Môi trường phát triển và triển khai',
    359: '7.2. Công nghệ xây dựng hệ thống',
    367: '7.3. Các công cụ OCR được tích hợp',
    374: '8. Xây dựng các thành phần chính của hệ thống',
    375: '8.1. Xây dựng backend, API và cơ sở dữ liệu',
    389: '8.2. Xây dựng giao diện người dùng',
    404: '8.3. Xây dựng chức năng quản trị hệ thống',
    414: '9. Xây dựng mô-đun xử lý tài liệu và OCR',
    415: '9.1. Tiếp nhận, chuẩn hóa và tiền xử lý tài liệu',
    426: '9.2. Tích hợp và lựa chọn công cụ OCR',
    437: '9.3. Thực hiện OCR và chuẩn hóa kết quả',
    449: '10. Triển khai và tối ưu hệ thống',
    450: '10.1. Triển khai hệ thống trong môi trường thử nghiệm',
    465: '10.2. Tối ưu hiệu năng xử lý và lưu trữ',
    473: '10.3. Hoàn thiện giao diện và trải nghiệm người dùng',
    484: '11. Thử nghiệm và kiểm thử chức năng hệ thống',
    485: '11.1. Môi trường thử nghiệm',
    495: '11.2. Bộ dữ liệu và phương pháp thử nghiệm',
    513: '11.3. Kiểm thử các chức năng người dùng',
    524: '11.4. Kiểm thử chức năng quản trị và phân quyền',
    531: '11.5. Tổng hợp kết quả kiểm thử chức năng',
    537: '12. Thực nghiệm và đánh giá chất lượng OCR',
    538: '12.1. Tiêu chí đánh giá',
    549: '12.2. Kết quả thử nghiệm trên các nhóm tài liệu',
    562: '12.3. So sánh các công cụ OCR',
    569: '12.4. Nhận xét và lựa chọn công cụ phù hợp',
    573: '13. Đánh giá hiệu năng và khả năng vận hành',
    574: '13.1. Thời gian xử lý và khả năng xử lý tài liệu nhiều trang',
    584: '13.2. Mức sử dụng tài nguyên và độ ổn định',
    590: '13.3. Đánh giá khả năng lưu trữ, truy vết và kiểm soát dữ liệu',
    597: '14. Nhận xét chung về kết quả xây dựng và thử nghiệm',
    598: '14.1. Kết quả đạt được',
    603: '14.2. Hạn chế và nguyên nhân',
    611: '14.3. Hướng hoàn thiện',
    77:  '2. Cơ sở thực tiễn',           # Chương 1: tiêu đề thiếu số thứ tự
}
for i, t in RENUM.items():
    set_par_text(ps[i], t)

d.save(DST)
print('đã lưu:', DST)
