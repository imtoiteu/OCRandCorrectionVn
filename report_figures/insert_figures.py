#!/usr/bin/env python3
"""Insert the generated report figures into the report, each directly above its
own "Hình X.Y." caption paragraph. Existing content is never removed."""
import re, sys, os, docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_IN, DOCX_OUT, PNG_DIR = sys.argv[1], sys.argv[2], sys.argv[3]
WIDTH = Inches(6.20)          # usable text width is 6.30 in

FIGURES = {
    '2.1':  'fig_2_1_use_case',                   '2.2':  'fig_2_2_context',
    '2.3':  'fig_2_3_architecture',               '2.4':  'fig_2_4_deployment',
    '2.5':  'fig_2_5_sequence_ocr',               '2.6':  'fig_2_6_document_workflow',
    '2.7':  'fig_2_7_engine_selection',           '2.8':  'fig_2_8_erd',
    '2.9':  'fig_2_9_navigation',                 '2.10': 'fig_2_10_data_protection',
    '3.1':  'fig_3_1_development_structure',      '3.2':  'fig_3_2_backend_modules',
    '3.6':  'fig_3_6_intake_normalization',       '3.7':  'fig_3_7_ocr_orchestration',
    '3.8':  'fig_3_8_ocr_result_representations', '3.9':  'fig_3_9_test_deployment',
    '4.1':  'fig_4_1_test_environment',           '4.2':  'fig_4_2_evaluation_procedure',
    '4.3':  'fig_4_3_functional_test_steps',      '4.4':  'fig_4_4_result_examples_template',
    '4.5':  'fig_4_5_processing_time_template',
}
# Deliberately not inserted: 3.3, 3.4, 3.5, 3.10, 3.11 (UI captures / no desktop app)

d = docx.Document(DOCX_IN)
cap_re = re.compile(r'^Hình\s+(\d+\.\d+)\.\s')

targets = {}
for p in d.paragraphs:
    m = cap_re.match(p.text.strip())
    if m:
        num = m.group(1)
        if num in targets:
            print('WARNING: duplicate caption for', num)
        targets[num] = p

missing = [n for n in FIGURES if n not in targets]
if missing:
    sys.exit('caption not found for: %s' % missing)

inserted = 0
for num, stem in sorted(FIGURES.items(), key=lambda kv: [int(x) for x in kv[0].split('.')]):
    png = os.path.join(PNG_DIR, stem + '.png')
    if not os.path.exists(png):
        sys.exit('missing PNG: ' + png)
    cap = targets[num]
    para = cap.insert_paragraph_before()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Inches(0.06)
    para.paragraph_format.space_after = Inches(0.04)
    para.add_run().add_picture(png, width=WIDTH)
    inserted += 1
    print('  Hình %-5s -> %s' % (num, stem + '.png'))

d.save(DOCX_OUT)
print('\ninserted %d figures; skipped 3.3, 3.4, 3.5, 3.10, 3.11 (manual captures)' % inserted)
